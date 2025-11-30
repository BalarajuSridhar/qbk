import os
import json
import zipfile
import re
from docx import Document
from docx.oxml.ns import qn

class DocxQuestionParser:
    def __init__(self, upload_folder):
        self.upload_folder = upload_folder
        self.images_folder = os.path.join(upload_folder, 'images')
        os.makedirs(self.images_folder, exist_ok=True)

    def extract_images_from_docx(self, docx_path):
        """Extract images physically stored in DOCX and save them in images folder"""
        images = {}
        try:
            with zipfile.ZipFile(docx_path, 'r') as docx_zip:
                for file_info in docx_zip.filelist:
                    if file_info.filename.startswith('word/media/'):
                        image_filename = os.path.basename(file_info.filename)
                        image_path = os.path.join(self.images_folder, image_filename)

                        with docx_zip.open(file_info.filename) as img:
                            with open(image_path, 'wb') as f:
                                f.write(img.read())

                        images[image_filename] = image_filename
                        print(f"📸 Extracted: {image_filename}")
        except Exception as e:
            print(f"❌ Error extracting images: {e}")

        return images

    def extract_images_from_cell(self, cell):
        """Extract image references from a specific table cell"""
        image_refs = []
        try:
            # Check for inline images in the cell
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    # Look for graphic data in the run
                    drawing = run._element.find('.//' + qn('w:drawing'))
                    if drawing is not None:
                        # Try to extract image reference
                        blip = drawing.find('.//' + qn('a:blip'))
                        if blip is not None:
                            embed = blip.get(qn('r:embed'))
                            if embed:
                                # This is an embedded image reference
                                image_ref = f"image{embed}.png"  # Simplified reference
                                image_refs.append(image_ref)
                                print(f"🖼 Found inline image in cell: {image_ref}")
        except Exception as e:
            print(f"⚠️ Error extracting images from cell: {e}")
        
        return image_refs

    def extract_image_references_from_text(self, text):
        """Extract image references from text using multiple pattern matching"""
        image_refs = []
        
        if not text:
            return image_refs
            
        # Pattern 1: Markdown style ![](media/imageX.png)
        pattern1 = r'!\[\]\(media/(image\d+\.png)\)'
        matches1 = re.findall(pattern1, text)
        
        # Pattern 2: HTML style with src attribute
        pattern2 = r'src=["\']media/(image\d+\.png)["\']'
        matches2 = re.findall(pattern2, text)
        
        # Pattern 3: Just the filename mentioned
        pattern3 = r'media/(image\d+\.png)'
        matches3 = re.findall(pattern3, text)
        
        # Pattern 4: Direct filename reference
        pattern4 = r'(image\d+\.png)'
        matches4 = re.findall(pattern4, text)
        
        # Combine all matches
        all_matches = matches1 + matches2 + matches3 + matches4
        
        for match in all_matches:
            if match not in image_refs:
                image_refs.append(match)
                print(f"🖼 Found image reference in text: {match}")
        
        return image_refs

    def is_hindi_text(self, text):
        """Check if text contains Hindi characters"""
        return bool(re.search(r'[\u0900-\u097F]', text))

    def clean_text(self, text):
        """Clean text by removing duplicates but keep original content"""
        if not text:
            return ""
        
        # Remove extra whitespace but keep the original text structure
        cleaned = re.sub(r'\s+', ' ', text).strip()
        
        # Remove exact duplicates (entire text repetition)
        if len(cleaned) > 10:
            mid = len(cleaned) // 2
            if cleaned[:mid].strip() == cleaned[mid:].strip():
                return cleaned[:mid].strip()
        
        return cleaned

    def question_has_image(self, question_data):
        """Check if a question actually has an image based on content analysis"""
        question_text = question_data.get("question_text", "").lower()
        solution_text = question_data.get("solution", "").lower()
        
        # Keywords that indicate the question likely has an image
        image_keywords = [
            # Direction questions
            'walks straight', 'walks towards', 'turns left', 'turns right', 'direction',
            'चलता है', 'मुड़ता है', 'दिशा',
            
            # Family relation questions  
            'brother of', 'daughter of', 'grandmother', 'defeated',
            'भाई', 'पुत्री', 'दादी', 'हराया',
            
            # Venn diagram questions
            'conclusions', 'statements', 'venn',
            'निष्कर्ष', 'कथन',
            
            # Pattern/visual questions
            'find the missing term', 'missing term', 'pattern',
            'लुप्त पद', 'आकृति',
            
            # Diagram questions
            'figure', 'diagram', 'chart', 'table',
            'चित्र', 'आरेख'
        ]
        
        # Check if question text contains any image-related keywords
        has_image_indicator = any(keyword in question_text for keyword in image_keywords)
        
        # Also check solution text for image references
        has_solution_image_ref = bool(self.extract_image_references_from_text(solution_text))
        
        return has_image_indicator or has_solution_image_ref

    def parse_docx(self, docx_path):
        """Main parse function - STORE ALL QUESTIONS SEPARATELY"""
        images = self.extract_images_from_docx(docx_path)
        document = Document(docx_path)

        questions = []
        image_keys = list(images.keys())
        image_usage_tracker = {img: False for img in image_keys}  # Track which images are used
        processed_questions = set()  # Track processed questions to avoid duplicates

        print(f"📸 Total images extracted: {len(image_keys)}")
        print(f"📸 Image files: {image_keys}")

        for table_index, table in enumerate(document.tables):
            print(f"\n🔍 Processing Table {table_index + 1}")
            
            # Check if this table contains English or Hindi
            table_language = self.detect_table_language(table)
            print(f"📝 Table language: {table_language}")
            
            if table_language == "english":
                # Process English question as separate question
                english_question = self.parse_question_table(table, "english", images, image_usage_tracker)
                if english_question and english_question.get("question_text"):
                    # Avoid duplicate questions
                    question_hash = hash(english_question["question_text"][:100])  # Use first 100 chars as hash
                    if question_hash not in processed_questions:
                        questions.append(english_question)
                        processed_questions.add(question_hash)
                        print(f"✅ Added English question: {english_question['question_text'][:80]}...")
                    else:
                        print(f"⚠️  Skipped duplicate English question")
            
            elif table_language == "hindi":
                # Process Hindi question as separate question
                hindi_question = self.parse_question_table(table, "hindi", images, image_usage_tracker)
                if hindi_question and hindi_question.get("question_text"):
                    # Avoid duplicate questions
                    question_hash = hash(hindi_question["question_text"][:100])  # Use first 100 chars as hash
                    if question_hash not in processed_questions:
                        questions.append(hindi_question)
                        processed_questions.add(question_hash)
                        print(f"✅ Added Hindi question: {hindi_question['question_text'][:80]}...")
                    else:
                        print(f"⚠️  Skipped duplicate Hindi question")

        # Print image usage summary
        used_images = [img for img, used in image_usage_tracker.items() if used]
        unused_images = [img for img, used in image_usage_tracker.items() if not used]
        
        print(f"\n🎉 Total Questions Parsed: {len(questions)}")
        print(f"📊 Language breakdown:")
        print(f"   English questions: {len([q for q in questions if not self.is_hindi_text(q.get('question_text', ''))])}")
        print(f"   Hindi questions: {len([q for q in questions if self.is_hindi_text(q.get('question_text', ''))])}")
        print(f"   Questions with images: {len([q for q in questions if q.get('image_path')])}")
        print(f"   Questions without images: {len([q for q in questions if not q.get('image_path')])}")
        print(f"📸 Image Usage:")
        print(f"   Used images: {len(used_images)}")
        print(f"   Unused images: {len(unused_images)}")
        
        return questions

    def detect_table_language(self, table):
        """Detect if table contains English or Hindi content"""
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if not any(cells):
                continue
            
            # Check first cell for language indicators
            first_cell = cells[0].lower()
            
            if "question" in first_cell or "type" in first_cell or "option" in first_cell:
                return "english"
            elif "प्रश्न" in first_cell or "प्रकार" in first_cell or "विकल्प" in first_cell:
                return "hindi"
            
            # Check content of other cells
            for cell_text in cells[1:]:
                if self.is_hindi_text(cell_text):
                    return "hindi"
                elif re.search(r'[a-zA-Z]', cell_text):
                    return "english"
        
        return "unknown"

    def parse_english_option_row(self, cell_texts, option_cell=None, images=None, image_usage_tracker=None):
        """Parse English option row with format: Option   text   correctness"""
        if len(cell_texts) < 2:
            return None
            
        option_text = cell_texts[1].strip()
        
        # For English options, correctness is usually in the 3rd column
        is_correct = False
        if len(cell_texts) >= 3:
            correctness_text = cell_texts[2].lower()
            is_correct = any(x in correctness_text for x in ["correct", "true", "right"])
        
        option_data = {
            "text": option_text,
            "is_correct": is_correct,
            "marks": 0,
            "image_path": None  # Initialize option image path
        }
        
        # Check if option cell has an image
        if option_cell and images and image_usage_tracker:
            option_images = self.extract_images_from_cell(option_cell)
            if option_images:
                for img_ref in option_images:
                    matching_images = [img for img in images 
                                     if (img_ref in img or img.startswith('image')) 
                                     and not image_usage_tracker.get(img, False)]
                    if matching_images:
                        option_data["image_path"] = matching_images[0]
                        image_usage_tracker[matching_images[0]] = True
                        print(f"🖼 Assigned option image: {matching_images[0]}")
                        break
        
        return option_data

    def parse_hindi_option_row(self, cell_texts, option_cell=None, images=None, image_usage_tracker=None):
        """Parse Hindi option row with format: विकल्प   text   correctness"""
        if len(cell_texts) < 2:
            return None
            
        option_text = cell_texts[1].strip()
        
        # For Hindi options, correctness is usually in the 3rd column
        is_correct = False
        if len(cell_texts) >= 3:
            correctness_text = cell_texts[2].lower()
            is_correct = any(x in correctness_text for x in ["सही", "ठीक", "हाँ"])
        
        option_data = {
            "text": option_text,
            "is_correct": is_correct,
            "marks": 0,
            "image_path": None  # Initialize option image path
        }
        
        # Check if option cell has an image
        if option_cell and images and image_usage_tracker:
            option_images = self.extract_images_from_cell(option_cell)
            if option_images:
                for img_ref in option_images:
                    matching_images = [img for img in images 
                                     if (img_ref in img or img.startswith('image')) 
                                     and not image_usage_tracker.get(img, False)]
                    if matching_images:
                        option_data["image_path"] = matching_images[0]
                        image_usage_tracker[matching_images[0]] = True
                        print(f"🖼 Assigned option image: {matching_images[0]}")
                        break
        
        return option_data

    def parse_question_table(self, table, language, images, image_usage_tracker):
        """Parse individual question table for either English or Hindi"""
        question_data = {
            "question_text": "",
            "type": "multiple_choice",
            "options": [],
            "correct_answer": "",
            "solution": "",
            "solution_image_path": None,  # New field for solution images
            "marks": 1,
            "image_path": None,
            "language": language
        }

        option_rows = []
        question_text_found = False
        image_keys = list(images.keys())

        for row_index, row in enumerate(table.rows):
            cells = [cell for cell in row.cells]
            cell_texts = [cell.text.strip() for cell in cells]
            
            if not any(cell_texts):
                continue

            key = cell_texts[0].lower()

            # Handle both English and Hindi headers
            if (("question" in key or "प्रश्न" in key) and not question_text_found):
                question_cell = cells[1] if len(cells) > 1 else None
                new_text = " ".join(cell_texts[1:]).strip()
                
                if new_text:
                    cleaned_text = self.clean_text(new_text)
                    question_data["question_text"] = cleaned_text
                    question_text_found = True
                    print(f"✅ {language.title()} question set: {cleaned_text[:80]}...")

                    # Check if question cell has an image
                    if question_cell:
                        question_cell_images = self.extract_images_from_cell(question_cell)
                        if question_cell_images:
                            for img_ref in question_cell_images:
                                matching_images = [img for img in image_keys 
                                                 if (img_ref in img or img.startswith('image')) 
                                                 and not image_usage_tracker.get(img, False)]
                                if matching_images:
                                    question_data["image_path"] = matching_images[0]
                                    image_usage_tracker[matching_images[0]] = True
                                    print(f"🖼 Assigned question cell image: {matching_images[0]}")
                                    break

            elif "type" in key or "प्रकार" in key:
                t = cell_texts[1].lower() if len(cell_texts) > 1 else ""
                if "integer" in t or "पूर्णांक" in t: 
                    question_data["type"] = "integer"
                elif "fill" in t or "रिक्त" in t: 
                    question_data["type"] = "fill_ups"
                elif "true" in t or "सत्य" in t or "false" in t or "असत्य" in t: 
                    question_data["type"] = "true_false"
                elif "comprehension" in t or "अवबोधन" in t: 
                    question_data["type"] = "comprehension"
                else: 
                    question_data["type"] = "multiple_choice"
                print(f"📂 Type: {question_data['type']}")

            elif "option" in key or "विकल्प" in key:
                # Parse options based on language
                option_cell = cells[1] if len(cells) > 1 else None
                if language == "english":
                    option_data = self.parse_english_option_row(cell_texts, option_cell, image_keys, image_usage_tracker)
                else:
                    option_data = self.parse_hindi_option_row(cell_texts, option_cell, image_keys, image_usage_tracker)
                
                if option_data and option_data["text"]:
                    option_rows.append(option_data)
                    image_status = " with image" if option_data.get("image_path") else ""
                    print(f"🔘 {language.title()} option: '{option_data['text']}' | Correct: {option_data['is_correct']}{image_status}")

            elif "answer" in key or "उत्तर" in key:
                if len(cell_texts) > 1:
                    question_data["correct_answer"] = cell_texts[1]
                    print(f"✔ Correct Answer from table: {cell_texts[1]}")

            elif "solution" in key or "उपाय" in key:
                solution_cell = cells[1] if len(cells) > 1 else None
                if len(cell_texts) > 1:
                    solution_text = " ".join(cell_texts[1:]).strip()
                    question_data["solution"] = solution_text
                    print(f"💡 Solution found for {language} question")
                    
                    # Check if solution cell has an image
                    if solution_cell:
                        solution_cell_images = self.extract_images_from_cell(solution_cell)
                        if solution_cell_images:
                            for img_ref in solution_cell_images:
                                matching_images = [img for img in image_keys 
                                                 if (img_ref in img or img.startswith('image')) 
                                                 and not image_usage_tracker.get(img, False)]
                                if matching_images:
                                    question_data["solution_image_path"] = matching_images[0]
                                    image_usage_tracker[matching_images[0]] = True
                                    print(f"🖼 Assigned solution image: {matching_images[0]}")
                                    break
                    
                    # Also check solution text for image references
                    if not question_data.get("solution_image_path"):
                        solution_image_refs = self.extract_image_references_from_text(solution_text)
                        if solution_image_refs:
                            for img_ref in solution_image_refs:
                                matching_images = [img for img in image_keys 
                                                 if img_ref in img 
                                                 and not image_usage_tracker.get(img, False)]
                                if matching_images:
                                    question_data["solution_image_path"] = matching_images[0]
                                    image_usage_tracker[matching_images[0]] = True
                                    print(f"🖼 Assigned solution text image: {matching_images[0]}")
                                    break

            elif "marks" in key or "अंक" in key:
                try:
                    if len(cell_texts) > 1:
                        marks_text = cell_texts[1]
                        marks_parts = marks_text.split()
                        if marks_parts:
                            question_data["marks"] = int(marks_parts[0])
                            print(f"📊 Correct Answer Marks: {question_data['marks']}")
                except:
                    question_data["marks"] = 1

        # Process options
        if question_data["type"] == "multiple_choice":
            self.process_multiple_choice_options(option_rows, question_data, language)
        else:
            self.process_other_options(option_rows, question_data, language)
        
        # Final image assignment check for question image
        if not question_data.get("image_path"):
            if self.question_has_image(question_data):
                unused_images = [img for img in image_keys if not image_usage_tracker.get(img, False)]
                if unused_images:
                    question_data["image_path"] = unused_images[0]
                    image_usage_tracker[unused_images[0]] = True
                    print(f"🖼 Content-based image assignment: {unused_images[0]}")
            else:
                print(f"ℹ️  No image assigned - question doesn't require one")
        
        return question_data

    def process_multiple_choice_options(self, option_rows, question_data, language):
        """Process options specifically for multiple_choice questions"""
        valid_options = []
        
        for option in option_rows:
            if option["text"] and option["text"].strip():
                valid_options.append(option)

        print(f"📦 Found {len(valid_options)} valid {language} options")
        
        # Ensure exactly 4 options for multiple choice
        if len(valid_options) >= 4:
            question_data["options"] = valid_options[:4]
        elif len(valid_options) > 0:
            question_data["options"] = valid_options
            # Add placeholder options
            for i in range(len(valid_options), 4):
                question_data["options"].append({
                    "text": f"Option {chr(65+i)}",
                    "is_correct": False,
                    "marks": 0,
                    "image_path": None
                })
        else:
            question_data["options"] = [
                {"text": "Option A", "is_correct": False, "marks": 0, "image_path": None},
                {"text": "Option B", "is_correct": False, "marks": 0, "image_path": None},
                {"text": "Option C", "is_correct": True, "marks": question_data["marks"], "image_path": None},
                {"text": "Option D", "is_correct": False, "marks": 0, "image_path": None}
            ]

        # Determine correct answer
        correct_options = [i for i, opt in enumerate(question_data["options"]) if opt["is_correct"]]
        if correct_options:
            correct_index = correct_options[0]
            question_data["correct_answer"] = chr(65 + correct_index)
            print(f"✅ Correct answer set to: {question_data['correct_answer']}")
        elif question_data.get("correct_answer"):
            print(f"✅ Using provided correct answer: {question_data['correct_answer']}")
        
        # Debug: Print all options
        for i, opt in enumerate(question_data["options"]):
            status = "✓" if opt["is_correct"] else "✗"
            image_status = " 📷" if opt.get("image_path") else ""
            print(f"  {chr(65+i)}. {opt['text']} [{status}]{image_status}")

    def process_other_options(self, option_rows, question_data, language):
        """Process options for non-multiple_choice questions"""
        valid_options = []
        
        for option in option_rows:
            if option["text"] and option["text"].strip():
                valid_options.append(option)

        question_data["options"] = valid_options
        
        correct_options = [i for i, opt in enumerate(valid_options) if opt["is_correct"]]
        if len(correct_options) == 1:
            question_data["correct_answer"] = chr(65 + correct_options[0])
            print(f"✅ Correct answer set to: {question_data['correct_answer']}")
        
        print(f"📦 {language} options: {len(question_data['options'])}")

# Usage example:
if __name__ == "__main__":
    parser = DocxQuestionParser("uploads")
    questions = parser.parse_docx("path/to/your/document.docx")
    
    for i, q in enumerate(questions):
        print(f"\n--- Question {i+1} ---")
        print(f"Text: {q['question_text'][:100]}...")
        print(f"Language: {q['language']}")
        print(f"Question Image: {q['image_path']}")
        print(f"Solution Image: {q.get('solution_image_path', 'None')}")
        print(f"Type: {q['type']}")
        print(f"Marks: {q['marks']}")
        print(f"Correct Answer: {q['correct_answer']}")
        print(f"Options: {len(q['options'])}")
        for j, opt in enumerate(q['options']):
            image_status = f" | Image: {opt.get('image_path', 'None')}" if opt.get('image_path') else ""
            print(f"  {chr(65+j)}. {opt['text']} (Correct: {opt['is_correct']}){image_status}")